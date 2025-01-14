from docx import Document

# Load the document
file_path = "demo.docx"  # Replace bằng đường dẫn file
doc = Document(file_path)

# Answer
answers = {
    # 1: 'A', 2: 'B', 3: 'D', 4: 'A', ... Gõ đáp án theo đúng format này
}


def bold_correct_answers(doc, answers):
    for idx, paragraph in enumerate(doc.paragraphs):
        # Process questions and answers in the document
        for question_num, correct_answer in answers.items():
            question_key = f"Câu {question_num}:"
            if question_key in paragraph.text:
                # Look for answers in subsequent paragraphs
                for i in range(1, 5):  # Assuming there are always 4 answer options (A, B, C, D)
                    if idx + i < len(doc.paragraphs):
                        answer_paragraph = doc.paragraphs[idx + i]
                        if answer_paragraph.text.strip().startswith(correct_answer + "."):
                            # Bold the entire line of the correct answer
                            for run in answer_paragraph.runs:
                                run.bold = True

# Apply the bolding function to the document
bold_correct_answers(doc, answers)

# Save the updated document
output_path = "demobold.docx"  # Output file name
doc.save(output_path)
print(f"Tài liệu đã được lưu: {output_path}")
